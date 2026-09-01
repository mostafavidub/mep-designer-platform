"""Generate the production MEP Design Rule Book v4 DOCX.

The runtime document mirrors the machine-enforced Drawing Manifest contract.
It is generated at deploy time so the deployed RULEBOOK_PATH is always the
same version used by the website proposal and Code Designer.
"""
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


VERSION = "4.1"
BENCHMARK = {
    "base_architectural_views": 4,
    "approved_deliverables": 29,
    "independent_issued_drawing_content": 29,
    "issued_layouts": 29,
}


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def build(path):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("EngiTools MEP Design Rule Book — Mechanical Drawing Set")
    r.bold = True; r.font.size = Pt(18)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Version {VERSION} — Approved Drawing Manifest Contract").bold = True

    add_heading(doc, "1. Source-of-truth contract")
    add_bullet(doc, "Architectural level count is not the mechanical deliverable count.")
    add_bullet(doc, "The customer-facing count is the number of separate approved mechanical drawings in the Drawing Manifest.")
    add_bullet(doc, "Proposal count = frozen Approved Drawing Manifest count = independent issued drawing-content count = issued CAD layout count.")
    add_bullet(doc, "Layout count alone is not proof that promised drawings were generated.")
    add_bullet(doc, "Any mismatch blocks issuance with: CAD output does not match approved drawing manifest.")

    add_heading(doc, "2. Authority drawing families")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, value in enumerate(("Code", "Family", "Typical content")):
        table.rows[0].cells[i].text = value
    rows = [
        ("M-W", "Water supply", "Cold/hot water, return, riser, equipment"),
        ("M-S", "Sanitary + vent", "Sanitary, vent, stacks, cleanouts, details"),
        ("M-H", "Heating", "Heating plans, equipment, supply/return"),
        ("M-C", "Cooling / HVAC", "Cooling, condensate, equipment, roof units"),
        ("M-G", "Gas", "Gas plans, meter/regulator, details"),
        ("M-V", "Ventilation / exhaust", "Exhaust, make-up air, parking/details"),
        ("M-R", "Roof / rainwater", "Roof drainage, drains and downpipes"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row): cells[i].text = value

    add_heading(doc, "3. First-class drawing types")
    add_bullet(doc, "FLOOR_PLAN, ROOF_PLAN, RISER_DIAGRAM, SCHEMATIC, EQUIPMENT_PLAN, DETAIL_SHEET, VENTILATION_PLAN are first-class manifest items.")
    add_bullet(doc, "A riser, equipment or detail requirement is not satisfied by renaming another floor-plan layout.")

    add_heading(doc, "4. Effective Levels and Typical Floors")
    add_bullet(doc, "Effective Levels are system-specific; each mechanical family receives only the levels where it applies.")
    add_bullet(doc, "Typical Floor consolidation is system-specific and conservative. A grouped Typical entry produces one issued drawing for that family only.")
    add_bullet(doc, "If fixture/equipment/routing evidence differs for a family, those levels remain separate even when architecture appears similar.")

    add_heading(doc, "5. Full-authority documentation trigger")
    add_bullet(doc, "Non-Typical multi-level projects with roof, vertical services and all principal mechanical families require the expanded authority package: independent riser/equipment/detail/schematic drawings in addition to system plans.")
    add_bullet(doc, "The trigger is based on project complexity and evidence, never a project name or hard-coded customer exception.")

    add_heading(doc, "6. Approved benchmark contract — 29 deliverables")
    p = doc.add_paragraph("For the approved benchmark architecture/reference set, the validated relationship is:")
    p.runs[0].bold = True
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
    table.rows[0].cells[0].text = "Measure"; table.rows[0].cells[1].text = "Required value"
    benchmark_rows = [
        ("Architectural base views", "4"),
        ("Approved mechanical deliverables", "29"),
        ("Independent issued drawing content", "29"),
        ("Issued CAD layouts", "29"),
        ("Parity", "29 = 29 = 29; architectural base views may differ"),
    ]
    for row in benchmark_rows:
        cells = table.add_row().cells; cells[0].text = row[0]; cells[1].text = row[1]
    add_bullet(doc, "The historical website value of 15 plans is not an accepted contract for this benchmark.")
    add_bullet(doc, "A file containing four architectural/model-space base plans may legitimately produce 29 authority deliverables, provided every manifest item has independent issued drawing content.")

    add_heading(doc, "7. Approval and immutability")
    add_bullet(doc, "The website shows the exact per-sheet Drawing Manifest before design starts: code, family, drawing type and represented level/Typical scope.")
    add_bullet(doc, "User approval freezes a content-hashed copy of the exact manifest. Post-approval proposal changes invalidate approval and require re-approval.")

    add_heading(doc, "8. Code Designer issuance")
    add_bullet(doc, "Code Designer iterates the frozen Approved Drawing Manifest; it must not recalculate or independently expand the sheet list.")
    add_bullet(doc, "Every standard system plan must contain the required system view; every special sheet must contain role-specific drawing geometry or a legitimate dedicated equipment/roof view.")
    add_bullet(doc, "Issued sheets receive EngiTools-owned dimensions, leaders and callouts after source drafting baggage is removed.")

    add_heading(doc, "9. Fail-closed QA")
    add_bullet(doc, "QA reports four different quantities: architectural base views, approved deliverables, independent issued drawing content, and issued layouts.")
    add_bullet(doc, "The required equality is approved deliverables = independent issued drawing content = issued layouts. Architectural base views are informational and may differ.")
    add_bullet(doc, "Missing/extra layouts, empty renamed sheets, missing role content, unresolved technical inputs or dirty DXF audit block issuance.")

    add_heading(doc, "10. Split-AC representation and visual acceptance")
    add_bullet(doc, "Every indoor unit shall use the standard ENGI_AC_INDOOR block with an IDU label, equipment tag, capacity callout, leader and airflow arrow.")
    add_bullet(doc, "Every outdoor unit shall use the standard ENGI_AC_OUTDOOR block with an ODU label, tag, served-IDU reference and refrigerant-riser callout.")
    add_bullet(doc, "IDU, ODU, refrigerant, condensate, callout, leader and schedule evidence shall remain semantically linked and count-consistent.")
    add_bullet(doc, "Each Split-AC sheet shall be rendered independently at release scale. A symbol below the minimum plotted pixel size or an empty preview blocks issuance.")
    add_bullet(doc, "Layer presence alone is not evidence of a visible or readable cooling unit.")

    add_heading(doc, "11. Professional responsibility")
    add_bullet(doc, "This automated package remains subject to professional engineering review and applicable statutory approval. The software does not claim statutory approval merely because automated QA passes.")

    doc.core_properties.title = "EngiTools MEP Design Rule Book v4"
    doc.core_properties.subject = "Mechanical approved drawing manifest and 29-deliverable benchmark"
    doc.save(path)
    print(f"installed Rule Book v4: {path} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else "/data/rulebook/MEP_Design_Rulebook.docx"
    build(target)
